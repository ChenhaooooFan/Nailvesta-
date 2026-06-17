"""
report_generator.py — Generate the NailVesta 中台运营周报 as a .docx file.

Produces the same section structure as the W24 report built in JS/docx,
but using python-docx so it can run inside Streamlit.
"""

from __future__ import annotations
import io, re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ─── COLOUR PALETTE ──────────────────────────────────────────────────────────
C = {
    "header": ("1F4E79", "FFFFFF"),   # dark blue bg, white text
    "text":    "333333",
    "good":    "1F6E3E",
    "bad":     "9C2B1B",
    "goodFill":"E2EFDA",
    "badFill": "FCE4D6",
    "warnFill":"FFF2CC",
    "noteFill":"D5E8F0",
    "altrow":  "F2F2F2",
    "neutral": "888888",
}


# ─── LOW-LEVEL HELPERS ───────────────────────────────────────────────────────

def _rgb(hex6: str) -> RGBColor:
    r, g, b = int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16)
    return RGBColor(r, g, b)


def _set_cell_bg(cell, hex6: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6.upper())
    tcPr.append(shd)


def _set_col_width(cell, twips: int):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old  = tcPr.find(qn("w:tcW"))
    if old is not None: tcPr.remove(old)
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.insert(0, tcW)


def _cell_para(cell, text: str | list, bold=False, color=None,
               size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    """Clear cell and write styled text. text can be str or list of (str,dict) tuples."""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align

    if isinstance(text, str):
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = _rgb(color)
    else:
        # list of (text, {bold, color, italic, size_pt})
        for segment in text:
            if isinstance(segment, str):
                run = p.add_run(segment)
                run.font.size = Pt(size_pt)
                run.font.color.rgb = _rgb(C["text"])
            else:
                s, opts = segment
                run = p.add_run(s)
                run.bold   = opts.get("bold", False)
                run.italic = opts.get("italic", False)
                run.font.size = Pt(opts.get("size_pt", size_pt))
                c = opts.get("color")
                if c:
                    run.font.color.rgb = _rgb(c)


def _table_borders(table):
    """Apply thin borders using python-docx's Table Grid style."""
    try:
        table.style = table._tbl.getroottree().getroot().find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styles"
        )
    except Exception:
        pass
    # Set borders via XML the OOXML-correct way (tblBorders after tblW)
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # Ensure tblW exists first (required before tblBorders per schema)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "0")
        tblW.set(qn("w:type"), "auto")
        tblPr.insert(0, tblW)
    # Remove old borders if present
    old_b = tblPr.find(qn("w:tblBorders"))
    if old_b is not None: tblPr.remove(old_b)
    borders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFBFBF")
        borders.append(el)
    # Insert tblBorders after tblW
    tblW_idx = list(tblPr).index(tblW)
    tblPr.insert(tblW_idx + 1, borders)


# ─── HIGH-LEVEL DOCUMENT HELPERS ─────────────────────────────────────────────

class ReportBuilder:
    """Wraps a python-docx Document with convenience methods."""

    def __init__(self):
        self.doc = Document()
        # Page: US Letter, 1.2 cm margins all sides
        sec  = self.doc.sections[0]
        sec.page_width  = Cm(21.59)
        sec.page_height = Cm(27.94)
        for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
            setattr(sec, attr, Cm(1.8))

        # Default font
        style = self.doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)
        style.font.color.rgb = _rgb(C["text"])

    def h1(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = _rgb("1F4E79")
        return p

    def h2(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = _rgb("2E75B6")
        return p

    def h3(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = _rgb("2E75B6")
        return p

    def para(self, text: str | list, size_pt=10, color=None,
             bold=False, italic=False, after_pt=4, before_pt=0):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before_pt)
        p.paragraph_format.space_after  = Pt(after_pt)
        if isinstance(text, str):
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            run.font.size = Pt(size_pt)
            if color: run.font.color.rgb = _rgb(color)
        else:
            for seg in text:
                if isinstance(seg, str):
                    r = p.add_run(seg)
                    r.font.size = Pt(size_pt)
                else:
                    s, opts = seg
                    r = p.add_run(s)
                    r.bold   = opts.get("bold", bold)
                    r.italic = opts.get("italic", italic)
                    r.font.size = Pt(opts.get("size_pt", size_pt))
                    c = opts.get("color")
                    if c: r.font.color.rgb = _rgb(c)
        return p

    def note(self, text: str, size_pt=8.5):
        """Small grey italic note paragraph."""
        return self.para(text, size_pt=size_pt, color=C["neutral"], italic=True)

    def insight_box(self, text: str | list, fill: str = "noteFill"):
        """Coloured insight/callout box implemented as a 1x1 table."""
        hex6 = C.get(fill, fill)
        tbl  = self.doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _table_borders(tbl)
        cell = tbl.rows[0].cells[0]
        _set_cell_bg(cell, hex6)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        if isinstance(text, str):
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.font.italic = True
            run.font.color.rgb = _rgb(C["text"])
        else:
            for seg in text:
                if isinstance(seg, str):
                    r = p.add_run(seg)
                    r.font.size = Pt(9.5)
                    r.font.italic = True
                else:
                    s, opts = seg
                    r = p.add_run(s)
                    r.bold   = opts.get("bold", False)
                    r.italic = opts.get("italic", True)
                    r.font.size = Pt(opts.get("size_pt", 9.5))
                    c = opts.get("color")
                    if c: r.font.color.rgb = _rgb(c)
        return tbl

    def table(self, headers: list[str], rows: list[list],
              col_widths_cm: list[float] | None = None,
              alt_rows: bool = True):
        """
        Build a styled table.
        Each cell value can be:
          - str  → plain text
          - (str, {bold, color, italic}) → styled text
        """
        ncols = len(headers)
        tbl   = self.doc.add_table(rows=1 + len(rows), cols=ncols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _table_borders(tbl)

        # Header row
        hrow = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            _set_cell_bg(cell, C["header"][0])
            if col_widths_cm:
                _set_col_width(cell, int(col_widths_cm[i] * 567))  # 1cm ≈ 567 twips
            _cell_para(cell, h, bold=True, color=C["header"][1], size_pt=9.5,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

        # Data rows
        for ri, row_data in enumerate(rows):
            drow = tbl.rows[ri + 1]
            bg   = C["altrow"] if (alt_rows and ri % 2 == 1) else "FFFFFF"
            for ci, val in enumerate(row_data):
                cell = drow.cells[ci]
                _set_cell_bg(cell, bg)
                if col_widths_cm:
                    _set_col_width(cell, int(col_widths_cm[ci] * 567))
                if isinstance(val, str):
                    _cell_para(cell, val, size_pt=9.5)
                elif isinstance(val, (int, float)):
                    _cell_para(cell, str(val), size_pt=9.5)
                elif isinstance(val, tuple) and len(val) == 2:
                    txt, opts = val
                    _cell_para(cell, str(txt),
                               bold=opts.get("bold", False),
                               color=opts.get("color"),
                               italic=opts.get("italic", False),
                               size_pt=9.5)
                else:
                    _cell_para(cell, str(val) if val is not None else "—", size_pt=9.5)

        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return tbl

    def page_break(self):
        self.doc.add_page_break()

    def spacer(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

    def to_bytes(self) -> bytes:
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()


# ─── FORMATTING HELPERS ──────────────────────────────────────────────────────

def _pct(v, decimals=1):
    if v is None: return "—"
    return f"{v*100:.{decimals}f}%" if v <= 1 else f"{v:.{decimals}f}%"

def _money(v, decimals=2):
    if v is None: return "—"
    return f"${v:,.{decimals}f}"

def _int(v):
    if v is None: return "—"
    return f"{int(v):,}"

def _wow(curr, prev, is_pct=False, invert=False):
    """Return (text, color) for WoW change. invert=True means lower is better."""
    if curr is None or prev is None: return "—", C["text"]
    diff = curr - prev
    if is_pct:
        txt = f"{'▲' if diff > 0 else '▼'} {abs(diff):.2f}pp"
    else:
        pct_chg = (curr - prev) / prev * 100 if prev else 0
        txt = f"{'▲' if diff > 0 else '▼'} {abs(pct_chg):.1f}%"
    good = diff < 0 if invert else diff > 0
    return txt, (C["good"] if good else C["bad"])

def _delta_pp(curr, prev):
    """pp delta for percentage values expressed as fractions."""
    if curr is None or prev is None: return "—", C["text"]
    d = (curr - prev) * 100
    txt = f"{'▲' if d > 0 else '▼'} {abs(d):.2f}pp"
    return txt, (C["good"] if d < 0 else C["bad"])


# ─── MAIN REPORT FUNCTION ────────────────────────────────────────────────────

def generate_report(
    week_num: int,
    date_range: str,
    order_metrics: dict,
    cancelled: dict,
    returned: dict,
    auction: dict,
    collection: dict,
    prev_week: dict | None = None,
    monthly_baseline: dict | None = None,  # last month average metrics
    may_baseline: dict | None = None,      # legacy, same as monthly_baseline if provided
    catalog_df=None,
) -> bytes:
    """
    Build the full weekly ops docx report and return as bytes.
    prev_week and may_baseline are optional; sections gracefully degrade.
    """
    m  = order_metrics
    pw = prev_week or {}
    # Monthly baseline: prefer explicit monthly_baseline, fall back to may_baseline
    mb_raw = monthly_baseline or may_baseline or {}
    mb = mb_raw  # shorthand
    # Month label for column headers
    mb_month = f"{mb_raw.get('year','上月')} 年 {mb_raw.get('month','')} 月均" if mb_raw.get("year") else "上月月均"
    mb_label = f"vs {mb_month}" if mb_raw else "vs 上月月均"

    rb = ReportBuilder()

    # ═══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════
    rb.para("")  # top margin
    rb.para("NailVesta", bold=True, size_pt=22, color="1F4E79",
            after_pt=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb.para("中台运营周报", bold=True, size_pt=18, color="1F4E79",
            after_pt=4).alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb.para(f"W{week_num} · {date_range}", bold=True, size_pt=13, color="2E75B6",
            after_pt=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb.para("取消 / 退货 / Auction / Collection 专题深度版 · 内部文件",
            size_pt=9, color=C["neutral"], italic=True,
            after_pt=16).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 数据范围提示 ──
    rb.insight_box(
        f"数据来源：全量订单 CSV + 4份专题 HTML 报告（Auction、取消订单、退货订单、Collection 链接综合分析）。"
        f"有效付费订单口径：全量 {_int(m['total_orders'])} − 0元单 {_int(m['zero_orders'])} "
        f"− 取消 {_int(m['cancelled_orders'])} = {_int(m['effective_orders'])} 单。"
        f"GMV 口径：SKU Subtotal After Discount（不含运费），与上周报告口径一致。",
        fill="warnFill"
    )
    rb.spacer()

    # ── 核心指标速览 ──
    rb.h2("核心指标速览")
    cr_txt, cr_col = _delta_pp(m["cancel_rate"], pw.get("cancel_rate"))
    rr_txt, rr_col = _delta_pp(m["return_rate"], pw.get("return_rate"))
    gmv_txt, gmv_col = _wow(m["gmv"], pw.get("gmv"))
    # Dynamic columns based on available baselines
    _mb_aov   = mb.get("aov") or (mb.get("gmv",0)/mb.get("effective_orders",1) if mb.get("effective_orders") else None)
    _mb_cr    = mb.get("cancel_rate")
    _mb_rr    = mb.get("return_rate")
    _mb_gmv_w = mb.get("weekly_avg",{}).get("gmv") or (mb.get("gmv",0)/mb.get("num_weeks",1) if mb.get("num_weeks") else mb.get("gmv"))

    if mb:
        rb.table(
            ["指标", f"本周 W{week_num}", f"W{week_num-1 if pw else '—'}", "WoW", mb_month, f"vs {mb_month.split('均')[0]}均"],
            [
                ["有效付费订单", _int(m["effective_orders"]), _int(pw.get("effective_orders")), _wow(m["effective_orders"], pw.get("effective_orders"))[0],
                 _int(mb.get("weekly_avg",{}).get("effective_orders") or mb.get("effective_orders")),
                 _wow(m["effective_orders"], mb.get("weekly_avg",{}).get("effective_orders") or mb.get("effective_orders"))[0]],
                ["0 元达人单", _int(m["zero_orders"]), "—", "—", "—", "—"],
                ["Cancelled 订单", _int(m["cancelled_orders"]), _int(pw.get("cancelled_orders")), _wow(m["cancelled_orders"], pw.get("cancelled_orders"), invert=True)[0],
                 _int(mb.get("weekly_avg",{}).get("cancelled_orders")), "—"],
                ["Cancel Rate", _pct(m["cancel_rate"]), _pct(pw.get("cancel_rate")), (cr_txt, {"color": cr_col, "bold": True}),
                 _pct(_mb_cr), (_delta_pp(m["cancel_rate"], _mb_cr)[0], {"color": _delta_pp(m["cancel_rate"], _mb_cr)[1], "bold": True})],
                ["Return Rate", _pct(m["return_rate"]), _pct(pw.get("return_rate")), (rr_txt, {"color": rr_col, "bold": True}),
                 _pct(_mb_rr), (_delta_pp(m["return_rate"], _mb_rr)[0], {"color": _delta_pp(m["return_rate"], _mb_rr)[1], "bold": True})],
                ["GMV（周均，不含运费）", _money(m["gmv"], 0), _money(pw.get("gmv"), 0), (gmv_txt, {"color": gmv_col}),
                 _money(_mb_gmv_w, 0), _wow(m["gmv"], _mb_gmv_w)[0]],
                ["AOV（不含运费）", _money(m["aov"]), _money(pw.get("aov")), _wow(m["aov"], pw.get("aov"))[0],
                 _money(_mb_aov), _wow(m["aov"], _mb_aov)[0]],
                ["SKU Sold", _int(m["sku_sold"]), _int(pw.get("sku_sold")), _wow(m["sku_sold"], pw.get("sku_sold"))[0], "—", "—"],
            ],
            col_widths_cm=[4.0, 2.5, 2.5, 2.5, 2.5, 3.0],
        )
    else:
        rb.table(
            ["指标", "本周 W" + str(week_num), "上周 W" + str(week_num-1) if pw else "上周", "WoW"],
            [
                ["有效付费订单", _int(m["effective_orders"]), _int(pw.get("effective_orders")), _wow(m["effective_orders"], pw.get("effective_orders"))[0]],
                ["0 元达人单", _int(m["zero_orders"]), "—", "—"],
                ["Cancelled 订单", _int(m["cancelled_orders"]), _int(pw.get("cancelled_orders")), _wow(m["cancelled_orders"], pw.get("cancelled_orders"), invert=True)[0]],
                ["Cancel Rate（付费口径）", _pct(m["cancel_rate"]), _pct(pw.get("cancel_rate")), (cr_txt, {"color": cr_col, "bold": True})],
                ["Return Rate（NailVesta 口径）", _pct(m["return_rate"]), _pct(pw.get("return_rate")), (rr_txt, {"color": rr_col, "bold": True})],
                ["GMV（SKU Sub，不含运费）", _money(m["gmv"], 0), _money(pw.get("gmv"), 0), (gmv_txt, {"color": gmv_col})],
                ["AOV（有效订单，不含运费）", _money(m["aov"]), _money(pw.get("aov")), _wow(m["aov"], pw.get("aov"))[0]],
                ["SKU Sold", _int(m["sku_sold"]), _int(pw.get("sku_sold")), _wow(m["sku_sold"], pw.get("sku_sold"))[0]],
                ["连带率（件/单）", f"{m['upo']:.2f}x" if m['upo'] else "—", f"{pw.get('upo'):.2f}x" if pw.get('upo') else "—", "—"],
            ],
            col_widths_cm=[5.5, 3.5, 3.5, 4.5],
        )

    rb.page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 一、核心销售指标
    # ═══════════════════════════════════════════════════════════════════════
    rb.h1("一、核心销售指标")
    rb.note(
        f"口径：有效付费订单 = 全量订单 {_int(m['total_orders'])} − 0元达人单 {_int(m['zero_orders'])} "
        f"− Cancelled {_int(m['cancelled_orders'])} = {_int(m['effective_orders'])}。"
        f"GMV = SKU Subtotal After Discount（不含运费）。Cancel Rate 分母 = 全量 − 0元单 = {_int(m['paid_base'])}。"
    )

    rb.h2("1.1 订单结构明细（含0元单与取消记录）")
    _mb_eff_w = mb.get("weekly_avg",{}).get("effective_orders") if mb else None
    _mb_can_w = mb.get("weekly_avg",{}).get("cancelled_orders") if mb else None
    _mb_cr_11 = mb.get("cancel_rate") if mb else None
    _mb_rr_11 = mb.get("return_rate") if mb else None
    rb.table(
        ["分层", f"本周 W{week_num}", "本周占全量%", f"W{week_num-1} 对比" if pw else "上周", f"{mb_month}（周均）" if mb else "上月月均"],
        [
            ["全量订单（Order ID 去重）", _int(m["total_orders"]), "100%", _int(pw.get("total_orders")), "—"],
            [("0元达人单（免费单）", {"color": C["neutral"]}), _int(m["zero_orders"]), _pct(m["zero_orders"]/m["total_orders"]), _int(pw.get("zero_orders")), "—"],
            [("Cancelled 订单（付费）", {"color": C["bad"]}),
             _int(m["cancelled_orders"]), _pct(m["cancelled_orders"]/m["total_orders"]),
             _int(pw.get("cancelled_orders")),
             f"~{_int(_mb_can_w)}/周" if _mb_can_w else "—"],
            [("有效付费订单", {"bold": True}),
             (_int(m["effective_orders"]), {"bold": True}),
             (_pct(m["effective_orders"]/m["total_orders"]), {"bold": True}),
             _int(pw.get("effective_orders")),
             f"~{_int(_mb_eff_w)}/周" if _mb_eff_w else "—"],
            [("Cancel Rate（付费口径）", {"bold": True}),
             (_pct(m["cancel_rate"]), {"bold": True, "color": C["good"] if (not _mb_cr_11 or m["cancel_rate"] < _mb_cr_11) else C["bad"]}),
             "—",
             _pct(pw.get("cancel_rate")),
             (_pct(_mb_cr_11), {"color": C["neutral"]}) if _mb_cr_11 else "—"],
            [("Return Rate（NailVesta）", {"bold": True}),
             (_pct(m["return_rate"]), {"bold": True, "color": C["good"] if (not _mb_rr_11 or m["return_rate"] < _mb_rr_11) else C["bad"]}),
             "—",
             _pct(pw.get("return_rate")),
             (_pct(_mb_rr_11), {"color": C["neutral"]}) if _mb_rr_11 else "—"],
            ["Cancelled SKU Units（件）", _int(m["cancelled_sku_units"]), "—", "—", "—"],
        ],
        col_widths_cm=[4.5, 2.2, 2.0, 2.5, 5.8],
    )

    rb.h2("1.2 GMV & 客单价")
    _mb_gmv_w2  = mb.get("weekly_avg",{}).get("gmv") if mb else None
    _mb_aov2    = mb.get("aov") if mb else None
    _mb_asp2    = mb.get("asp") if mb else None
    if mb:
        rb.table(
            ["指标", "本周", "上周 WoW", "WoW变动", f"{mb_month}（周均）", f"vs {mb_month.split('均')[0]}均"],
            [
                [("GMV（SKU Sub，不含运费）", {"bold": True}), (_money(m["gmv"], 0), {"bold": True}),
                 _money(pw.get("gmv"), 0) if pw else "—", _wow(m["gmv"], pw.get("gmv"))[0] if pw else "—",
                 _money(_mb_gmv_w2, 0) if _mb_gmv_w2 else "—", _wow(m["gmv"], _mb_gmv_w2)[0] if _mb_gmv_w2 else "—"],
                ["AOV（不含运费）", _money(m["aov"]), _money(pw.get("aov")) if pw else "—",
                 _wow(m["aov"], pw.get("aov"))[0] if pw else "—",
                 _money(_mb_aov2) if _mb_aov2 else "—", _wow(m["aov"], _mb_aov2)[0] if _mb_aov2 else "—"],
                ["ASP 件单价", _money(m["asp"]), _money(pw.get("asp")) if pw else "—",
                 _wow(m["asp"], pw.get("asp"))[0] if pw else "—",
                 _money(_mb_asp2) if _mb_asp2 else "—", _wow(m["asp"], _mb_asp2)[0] if _mb_asp2 else "—"],
                ["Order Amount（含运费，参考）", _money(m["order_amount_total"], 0), "—", "—", "—", "不同口径"],
                ["AOV（含运费，参考）", _money(m["aov_incl_shipping"]), "—", "—", "—", "不同口径"],
            ],
            col_widths_cm=[4.0, 2.5, 2.5, 2.2, 2.5, 3.3],
        )
    else:
        rb.table(
            ["指标", "本周", "上周", "WoW"],
            [
                [("GMV（SKU Sub After Discount）", {"bold": True}), (_money(m["gmv"], 0), {"bold": True}), _money(pw.get("gmv"), 0), _wow(m["gmv"], pw.get("gmv"))[0] if pw else "—"],
                ["AOV（有效订单，不含运费）", _money(m["aov"]), _money(pw.get("aov")), _wow(m["aov"], pw.get("aov"))[0] if pw else "—"],
                ["ASP 件单价（GMV÷SKU Sold）", _money(m["asp"]), _money(pw.get("asp")), _wow(m["asp"], pw.get("asp"))[0] if pw else "—"],
                ["Order Amount（含运费，仅参考）", _money(m["order_amount_total"], 0), "—", "不含运费口径见上方 GMV"],
                ["AOV（含运费）", _money(m["aov_incl_shipping"]), "—", "与 GMV 口径不同，仅参考"],
            ],
            col_widths_cm=[6, 3, 3, 5],
        )
    _insight_parts = []
    if pw.get("gmv"):
        gmv_d = (m["gmv"] - pw["gmv"]) / pw["gmv"] * 100
        _insight_parts.append(f"WoW GMV {'▲' if gmv_d>0 else '▼'}{abs(gmv_d):.1f}%（{_money(pw['gmv'],0)}→{_money(m['gmv'],0)}）")
    if _mb_gmv_w2:
        gmv_mom = (m["gmv"] - _mb_gmv_w2) / _mb_gmv_w2 * 100
        _insight_parts.append(f"vs {mb_month} {'▲' if gmv_mom>0 else '▼'}{abs(gmv_mom):.1f}%（周均{_money(_mb_gmv_w2,0)}）")
    if _insight_parts:
        rb.insight_box("  ·  ".join(_insight_parts), fill="goodFill" if (not pw.get("gmv") or m["gmv"] > pw["gmv"]) else "badFill")

    rb.h2("1.3 件数结构与 AOV 分层")
    if mb:
        rb.table(
            ["件数", "本周占比", "本周 AOV", "上周占比", "上周 AOV", f"{mb_month}占比", f"{mb_month}AOV"],
            [
                [label,
                 _pct(d["pct"]), _money(d["aov"]),
                 _pct(pw.get("qty_bands", {}).get(label, {}).get("pct")) if pw else "—",
                 _money(pw.get("qty_bands", {}).get(label, {}).get("aov")) if pw else "—",
                 _pct(mb.get("qty_bands", {}).get(label, {}).get("pct")),
                 _money(mb.get("qty_bands", {}).get(label, {}).get("aov")),
                ]
                for label, d in m["qty_bands"].items()
            ],
            col_widths_cm=[2.0, 2.0, 2.3, 2.0, 2.3, 2.2, 4.2],
        )
    else:
        rb.table(
            ["件数", "本周订单数", "本周占比", "本周 AOV", "上周占比", "上周 AOV", "AOV WoW"],
            [
                [label,
                 _int(d["count"]), _pct(d["pct"]), _money(d["aov"]),
                 _pct(pw.get("qty_bands", {}).get(label, {}).get("pct")),
                 _money(pw.get("qty_bands", {}).get(label, {}).get("aov")),
                 _wow(d["aov"], pw.get("qty_bands", {}).get(label, {}).get("aov"))[0] if pw else "—",
                ]
                for label, d in m["qty_bands"].items()
            ],
            col_widths_cm=[2, 2, 2, 2.5, 2, 2.5, 4],
        )

    rb.h2("1.4 支付方式分布")
    total_pm = sum(m["payment_dist"].values())
    pm_rows = []
    for method, cnt in sorted(m["payment_dist"].items(), key=lambda x: -x[1]):
        pct  = cnt / total_pm
        note = ""
        if "PayLater" in method or "Klarna" in method or "Affirm" in method or "over time" in method:
            note = ("先买后付", {"color": C["bad"]})
        pm_rows.append([method, _int(cnt), _pct(pct), note or ""])
    rb.table(["支付方式", "订单数", "占比", "备注"], pm_rows, col_widths_cm=[6, 2, 2, 7])

    rb.h2("1.5 AOV 分布")
    aov_labels = ["<$20","$20-30","$30-40","$40-60","$60-80","$80-120","$120+"]
    rb.table(
        ["AOV 区间", "本周单数", "本周占比", "上周占比", "WoW"],
        [
            [lbl,
             _int(m["aov_dist"].get(lbl, 0)),
             _pct(m["aov_dist"].get(lbl, 0) / m["effective_orders"]) if m["effective_orders"] else "—",
             _pct(pw.get("aov_dist", {}).get(lbl, 0) / pw.get("effective_orders", 1)) if pw.get("effective_orders") else "—",
             (lambda c, p: _delta_pp(c, p))(
                 m["aov_dist"].get(lbl, 0) / m["effective_orders"] if m["effective_orders"] else None,
                 pw.get("aov_dist", {}).get(lbl, 0) / pw.get("effective_orders", 1) if pw.get("effective_orders") else None
             )[0] if pw else "—",
             ]
            for lbl in aov_labels
        ],
        col_widths_cm=[2.5, 2.5, 2.5, 2.5, 7],
    )

    rb.h2("1.6 原价区间 GMV（SKU Sub After Discount）")
    price_labels = ["≤$29.99","$34.99","$39.99","$44.99","$49.99","$54.99","$55+"]
    total_pgmv = m["total_price_gmv"] or 1
    rb.table(
        ["原价区间", "本周 GMV", "本周占比", "上周 GMV", "上周占比", "占比变动"],
        [
            [lbl,
             _money(m["price_gmv"].get(lbl, 0), 0),
             _pct(m["price_gmv"].get(lbl, 0) / total_pgmv),
             _money(pw.get("price_gmv", {}).get(lbl, 0), 0) if pw else "—",
             _pct(pw.get("price_gmv", {}).get(lbl, 0) / pw.get("total_price_gmv", 1)) if pw.get("total_price_gmv") else "—",
             (lambda c,p: _delta_pp(c,p))(
                 m["price_gmv"].get(lbl,0)/total_pgmv,
                 pw.get("price_gmv",{}).get(lbl,0)/pw.get("total_price_gmv",1) if pw.get("total_price_gmv") else None
             )[0] if pw else "—",
             ]
            for lbl in price_labels
        ],
        col_widths_cm=[2.5, 2.5, 2, 2.5, 2, 5.5],
    )

    rb.page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 二、取消分析
    # ═══════════════════════════════════════════════════════════════════════
    rb.h1("二、取消分析")
    rb.note(
        "数据来自取消订单 HTML 分析报告。总订单量为 HTML 报告中的全量数字（含 0元单），"
        "Cancel Rate 以 CSV 付费口径为准（见第一节）。"
    )

    cr = cancelled
    rb.h2("2.1 直播 vs 非直播")
    live_rows = []
    if cr.get("live_sessions"):
        for s in cr["live_sessions"]:
            live_rows.append([s["name"], _int(s.get("cancelled")), _int(s.get("total")),
                              _pct(s.get("rate_pct"), 2) if s.get("rate_pct") else "—",
                              _pct(s.get("share_pct"), 2) if s.get("share_pct") else "—"])
    if cr.get("live_cancelled") is not None:
        live_rows.append([("直播合计", {"bold": True}), (_int(cr["live_cancelled"]), {"bold": True}),
                          _int(cr.get("live_total_orders")),
                          (_pct(cr.get("live_cancel_rate_pct"), 2) if cr.get("live_cancel_rate_pct") else "—", {"bold": True}), ""])
    if cr.get("non_live_cancelled") is not None:
        live_rows.append(["非直播", _int(cr["non_live_cancelled"]), _int(cr.get("non_live_total_orders")),
                          _pct(cr.get("non_live_cancel_rate_pct"), 2) if cr.get("non_live_cancel_rate_pct") else "—", ""])
    if live_rows:
        rb.table(["细分", "取消单数", "总订单量", "取消率", "占总取消%"],
                 live_rows, col_widths_cm=[4, 2.5, 3, 2.5, 5])

    rb.h2("2.2 取消时段分布（24小时）")
    if cr.get("hourly_counts") and len(cr["hourly_counts"]) == 24:
        hc = cr["hourly_counts"]
        total_h = sum(hc)
        blocks = {
            "凌晨 0-5": sum(hc[0:6]),
            "上午 6-11": sum(hc[6:12]),
            "午后 12-17": sum(hc[12:18]),
            "晚间 18-23": sum(hc[18:24]),
        }
        rb.table(
            ["时段", "取消事件数", "占比"],
            [[k, _int(v), _pct(v/total_h) if total_h else "—"] for k, v in blocks.items()],
            col_widths_cm=[4, 3, 10],
        )
        peak_h = hc.index(max(hc))
        rb.insight_box(f"全周取消峰值时段：{peak_h}点（{hc[peak_h]}件，{hc[peak_h]/total_h*100:.1f}%）", fill="warnFill")

    rb.h2("2.3 取消原因")
    if cr.get("cancel_reasons"):
        rb.table(
            ["取消原因", "单数", "占比"],
            [[r["reason"], _int(r["count"]), _pct(r.get("pct"), 1) if r.get("pct") else "—"]
             for r in cr["cancel_reasons"]],
            col_widths_cm=[8, 2.5, 6.5],
        )

    rb.h2("2.4 Collection 链接取消分布")
    if cr.get("collection_cancel"):
        rb.table(
            ["Collection 链接", "Cancelled", "占比"],
            [[r["collection"], _int(r["cancelled"]), _pct(r.get("share_pct"), 1) if r.get("share_pct") else "—"]
             for r in cr["collection_cancel"][:12]],
            col_widths_cm=[8, 2.5, 6.5],
        )

    rb.page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 三、退货分析
    # ═══════════════════════════════════════════════════════════════════════
    rb.h1("三、退货分析")
    rb.note(
        "退货率口径（NailVesta）：Return Rate = (Items Canceled + Items Returned) ÷ Items Sold。"
        "Items = SKU 件数（行级别）；Items Sold = 所有付费订单 Quantity 合计（含取消）；"
        "Items Returned = 退货报告行总数（每行 = 1 件）。"
    )
    ret = returned

    rb.h2("3.1 核心指标")
    rb.table(
        ["指标", "本周", "上周", "WoW"],
        [
            [("Return Rate（NailVesta 口径）", {"bold": True}),
             (_pct(m["return_rate"]), {"bold": True, "color": C["bad"] if m["return_rate"] and m["return_rate"] > 0.08 else C["good"]}),
             _pct(pw.get("return_rate")),
             (lambda c,p: _delta_pp(c,p))(m["return_rate"], pw.get("return_rate"))[0] if pw else "—"],
            ["Items Sold（付费订单件数合计）", _int(m["items_sold"]), _int(pw.get("items_sold")), "—"],
            ["Items Canceled（取消件数）", _int(m["items_canceled"]), _int(pw.get("items_canceled")), "—"],
            ["Items Returned（退货报告行总数）", _int(m["items_returned"]), _int(pw.get("items_returned")), "—"],
            ["退货去重包裹数", _int(ret.get("deduped_packages")), _int(pw.get("deduped_packages")), "—"],
        ],
        col_widths_cm=[6, 3, 3, 5],
    )

    rb.h2("3.2 退货构成（行级标签）")
    if any(ret.get(k) is not None for k in ["seller_fault_rows","refund_only","request_cancelled","shipped_return"]):
        total_ret = ret.get("total_rows") or m["items_returned"]
        def _rpct(v):
            return _pct(v / total_ret) if (v and total_ret) else "—"
        rb.table(
            ["分类", "本周行数", "占比（/退货行总数）", "WoW"],
            [
                ["Seller Fault（平台标注）", _int(ret.get("seller_fault_rows")), _rpct(ret.get("seller_fault_rows")), "—"],
                ["Refund Only", _int(ret.get("refund_only")), _rpct(ret.get("refund_only")), "—"],
                ["Request Cancelled", _int(ret.get("request_cancelled")), _rpct(ret.get("request_cancelled")), "—"],
                ["已寄出退回", _int(ret.get("shipped_return")), _rpct(ret.get("shipped_return")), "—"],
            ],
            col_widths_cm=[5, 2.5, 4, 5.5],
        )

    rb.h2("3.3 退货原因 Top 10")
    if ret.get("return_reasons"):
        rb.table(
            ["退货原因", "行数", "占比"],
            [[r["reason"], _int(r["count"]), _pct(r.get("pct"), 1) if r.get("pct") else "—"]
             for r in ret["return_reasons"][:10]],
            col_widths_cm=[8, 2.5, 6.5],
        )

    rb.h2("3.4 高退货款式 Top 10")
    if ret.get("style_returns"):
        rb.table(
            ["款式", "退货行数", "占比"],
            [[r["style"], _int(r["count"]), _pct(r.get("pct"), 1) if r.get("pct") else "—"]
             for r in ret["style_returns"][:10]],
            col_widths_cm=[8, 2.5, 6.5],
        )

    rb.h2("3.5 高退货产品链接")
    if ret.get("collection_returns"):
        rb.table(
            ["产品链接", "退货行数", "占比"],
            [[r["collection"], _int(r["rows"]), _pct(r.get("pct"), 1) if r.get("pct") else "—"]
             for r in ret["collection_returns"][:12]],
            col_widths_cm=[8, 2.5, 6.5],
        )

    rb.page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 四、Auction 专线
    # ═══════════════════════════════════════════════════════════════════════
    rb.h1("四、Auction 专线")
    auc = auction
    rb.table(
        ["指标", "本周", "上周", "WoW"],
        [
            ["总订单量（全量）", _int(auc.get("total_orders")), "—", "—"],
            ["取消订单", _int(auc.get("cancelled")), "—", "—"],
            [("Cancel Rate", {"bold": True}), (_pct(auc.get("cancel_rate_pct"), 1) if auc.get("cancel_rate_pct") else "—", {"bold": True}), "—", "—"],
            ["有效订单", _int(auc.get("effective_orders")), "—", "—"],
            ["有效 AOV", _money(auc.get("aov")), "—", "—"],
            ["退货申请", _int(auc.get("returns")), "—", "—"],
            ["退货率", _pct(auc.get("return_rate_pct"), 1) if auc.get("return_rate_pct") else "—", "—", "—"],
            ["取消原因构成", auc.get("cancel_reason", "—"), "—", "—"],
        ],
        col_widths_cm=[5, 3, 3, 6],
    )

    rb.page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 五、Collection 渠道综合分析
    # ═══════════════════════════════════════════════════════════════════════
    rb.h1("五、Collection 渠道综合分析")
    coll = collection

    rb.h2("5.1 链接明细：退货 vs 取消")
    if coll.get("links"):
        rb.table(
            ["Collection 链接", "渠道", "退货行数", "退货占比", "Cancelled", "Cancel占比", "差值 pp"],
            [
                [lk["collection"], lk.get("channel","—"),
                 _int(lk.get("return_rows")), _pct(lk.get("return_pct"), 1) if lk.get("return_pct") else "—",
                 _int(lk.get("cancelled")), _pct(lk.get("cancel_pct"), 1) if lk.get("cancel_pct") else "—",
                 f"{lk.get('diff_pp',0):+.2f}" if lk.get("diff_pp") is not None else "—"]
                for lk in coll["links"]
            ],
            col_widths_cm=[3.5, 2, 2, 2, 2, 2, 3.5],
        )

    rb.h2("5.2 渠道层级汇总")
    if coll.get("channel_summary"):
        rb.table(
            ["渠道类型", "退货行数", "退货占比", "Cancelled", "Cancel占比"],
            [
                [ch["channel"], _int(ch.get("return_rows")),
                 _pct(ch.get("return_pct"), 1) if ch.get("return_pct") else "—",
                 _int(ch.get("cancelled")), _pct(ch.get("cancel_pct"), 1) if ch.get("cancel_pct") else "—"]
                for ch in coll["channel_summary"]
            ],
            col_widths_cm=[4, 2.5, 3, 2.5, 5],
        )

    rb.page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 六、数据范围与口径说明
    # ═══════════════════════════════════════════════════════════════════════
    rb.h1("六、数据范围与口径说明")
    rb.table(
        ["口径", "本周定义"],
        [
            ["有效付费订单", f"全量 {_int(m['total_orders'])} − 0元单 {_int(m['zero_orders'])} − Cancelled {_int(m['cancelled_orders'])} = {_int(m['effective_orders'])}"],
            ["Cancel Rate", f"Cancelled / Paid Base = {_int(m['cancelled_orders'])} / {_int(m['paid_base'])} = {_pct(m['cancel_rate'])}"],
            ["Return Rate", f"(Items Canceled {_int(m['items_canceled'])} + Items Returned {_int(m['items_returned'])}) / Items Sold {_int(m['items_sold'])} = {_pct(m['return_rate'])}"],
            ["GMV", "SKU Subtotal After Discount（折后，不含运费）"],
            ["AOV", "GMV ÷ 有效付费订单数（不含运费）"],
            ["Items Sold", "所有付费订单（Amount>0）的 Quantity 合计，含已取消订单件数"],
            ["0元达人单", "Order Amount = 0，Payment Method = 空，不计入任何销售指标"],
        ],
        col_widths_cm=[4, 13],
    )

    rb.page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 七、综合结论与行动建议
    # ═══════════════════════════════════════════════════════════════════════
    rb.h1("七、综合结论与行动建议")

    # Auto-generate P0/P1 based on metrics
    p0_items = []
    p1_items = []

    if m["cancel_rate"] and m["cancel_rate"] > 0.04:
        p0_items.append(f"Cancel Rate {_pct(m['cancel_rate'])} 偏高（>4%），需排查直播取消来源")
    if m["return_rate"] and m["return_rate"] > 0.12:
        p0_items.append(f"Return Rate {_pct(m['return_rate'])} 偏高（>12%），需排查 Seller Fault 链路")

    if ret.get("seller_fault_rows") and ret.get("total_rows"):
        sf_pct = ret["seller_fault_rows"] / (ret["total_rows"] or 1)
        if sf_pct > 0.35:
            p0_items.append(f"Seller Fault 占比 {sf_pct*100:.1f}% 超过 35%，需仓库/质检/供应链分环节整改")

    if coll.get("links"):
        top_cancel = max(coll["links"], key=lambda x: x.get("cancel_pct") or 0, default=None)
        if top_cancel and top_cancel.get("cancel_pct", 0) > 15:
            p0_items.append(f"{top_cancel['collection']} Cancel 占比 {top_cancel.get('cancel_pct', 0):.1f}%，建议专项排查")

    if auc.get("cancel_rate_pct") and auc.get("cancel_rate_pct") > 8:
        p1_items.append(f"Auction Cancel Rate {_pct(auc.get('cancel_rate_pct'), 1)}（目标≤8%），需加强付款提醒")

    if ret.get("shipped_return"):
        p1_items.append(f"已寄出退回 {_int(ret.get('shipped_return'))} 件，建议复核仓库退货处理时效")

    rb.h2("7.1 P0 优先行动项")
    if p0_items:
        for item in p0_items:
            p = rb.doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            run.font.size = Pt(10)
            run.font.color.rgb = _rgb(C["bad"])
    else:
        rb.para("本周暂无明显 P0 风险", color=C["good"])

    rb.h2("7.2 P1 关注项")
    if p1_items:
        for item in p1_items:
            p = rb.doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            run.font.size = Pt(10)
    else:
        rb.para("暂无额外 P1 事项")

    rb.insight_box(
        "如需完整行动建议（含话术、跨部门分工、两周改善目标），请结合本报告数据在 Claude 中发起深度分析。",
        fill="noteFill"
    )

    return rb.to_bytes()
